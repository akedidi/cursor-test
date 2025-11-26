import os
import glob
import csv
import statistics
import re
import logging
from collections import defaultdict
from datetime import datetime

from dotenv import load_dotenv
import xlsxwriter

from docx import Document
from docx.oxml import OxmlElement
from docx.table import Table


# --------------------------------------------------------
# Logging
# --------------------------------------------------------
logging.basicConfig(
    level=logging.INFO,
    format="%(asctime)s [%(levelname)s] %(message)s",
)

# Ordre des labels qu'on veut dans les synthèses
LABEL_ORDER = [
    "Genera Token",   # Token
    "Purchase",
    "Policy",
    "Generate PDF",
    "Cancel",
]


# --------------------------------------------------------
# .env
# --------------------------------------------------------
def load_env():
    load_dotenv()

    results_folder = os.getenv("RESULTS_FOLDER")
    output_file = os.getenv("OUTPUT_FILE", "recap_scenarios.xlsx")
    doc_template = os.getenv("DOC_TEMPLATE")
    doc_output = os.getenv("DOC_OUTPUT")

    logging.info("RESULTS_FOLDER = %s", results_folder)
    logging.info("OUTPUT_FILE   = %s", output_file)
    logging.info("DOC_TEMPLATE  = %s", doc_template)
    logging.info("DOC_OUTPUT    = %s", doc_output)

    if not results_folder:
        raise ValueError("La variable RESULTS_FOLDER n'est pas définie dans le fichier .env")
    if not os.path.isdir(results_folder):
        raise ValueError(f"Le dossier RESULTS_FOLDER n'existe pas : {results_folder}")

    # si OUTPUT_FILE est un dossier ou sans extension -> on ajoute un nom
    if os.path.isdir(output_file) or not os.path.splitext(output_file)[1]:
        output_file = os.path.join(output_file, "recap_scenarios.xlsx")
        logging.info("OUTPUT_FILE normalisé en : %s", output_file)

    return results_folder, output_file, doc_template, doc_output


# --------------------------------------------------------
# Recherche des fichiers
# --------------------------------------------------------
def extract_users_from_filename(path: str) -> int:
    """
    Extrait le nombre d'utilisateurs à partir du nom de fichier.
    Ex : ...results-1-users.csv -> 1
         ...results-12-user.csv -> 12
    """
    name = os.path.basename(path)
    m = re.search(r"results-(\d+)-user", name)
    if m:
        return int(m.group(1))
    return 999999  # au cas où


def find_scenario_files(results_folder: str):
    """
    On accepte :
      IDP API-results-1-user.csv
      IDP API-results-1-users.csv
    """
    pattern = os.path.join(results_folder, "IDP API-results-*user*.csv")
    logging.info("Recherche des fichiers avec le pattern : %s", pattern)
    files = glob.glob(pattern)

    if not files:
        raise FileNotFoundError(f"Aucun fichier trouvé avec le pattern : {pattern}")

    # Tri par nombre d'utilisateurs : 1,2,4,8,12...
    files = sorted(files, key=extract_users_from_filename)

    logging.info("Nombre de fichiers trouvés : %d", len(files))
    for f in files:
        logging.info(" - %s", f)

    return files


# --------------------------------------------------------
# Lecture CSV
# --------------------------------------------------------
def read_jmeter_csv(path: str):
    logging.info("Lecture du fichier CSV : %s", path)
    rows = []
    with open(path, newline="", encoding="utf-8") as f:
        reader = csv.DictReader(f)
        for r in reader:
            rows.append(r)
    logging.info("  -> %d lignes lues (hors en-tête)", len(rows))
    return rows


# --------------------------------------------------------
# Helpers
# --------------------------------------------------------
def to_float(value, default=None):
    try:
        return float(value)
    except Exception:
        return default


def to_bool_success(value):
    if value is None:
        return False
    v = str(value).strip().lower()
    return v in ("true", "1", "yes", "y")


def percentile(values, p):
    """
    Percentile simple (0-100) sur une liste de nombres
    """
    if not values:
        return None
    values = sorted(values)
    k = (len(values) - 1) * (p / 100.0)
    f = int(k)
    c = min(f + 1, len(values) - 1)
    if f == c:
        return values[int(k)]
    d0 = values[f] * (c - k)
    d1 = values[c] * (k - f)
    return d0 + d1


def compute_execution_range_string(rows):
    """
    Calcule la date/heure de début et fin du scénario à partir des timeStamp JMeter.
    Format : 21/11/25 09:42 PM - 21/11/25 10:00 PM
    """
    timestamps = []
    for r in rows:
        ts = r.get("timeStamp")
        if ts is None:
            continue
        try:
            timestamps.append(int(ts))
        except ValueError:
            continue

    if not timestamps:
        return ""

    start_ts = min(timestamps) / 1000.0
    end_ts = max(timestamps) / 1000.0

    start_dt = datetime.fromtimestamp(start_ts)
    end_dt = datetime.fromtimestamp(end_ts)

    fmt = "%d/%m/%y %I:%M %p"  # 21/11/25 09:42 PM
    return f"{start_dt.strftime(fmt)} - {end_dt.strftime(fmt)}"


# --------------------------------------------------------
# Calcul du recap par scénario
# --------------------------------------------------------
def compute_recap(rows):
    logging.info("Calcul du tableau récapitulatif pour le scénario...")
    labels = {}

    for r in rows:
        label = r.get("label")
        elapsed_raw = r.get("elapsed")
        success_raw = r.get("success")

        if label is None or elapsed_raw is None:
            continue

        elapsed = to_float(elapsed_raw)
        if elapsed is None:
            continue

        success = to_bool_success(success_raw)

        if label not in labels:
            labels[label] = {"times": [], "errors": 0}

        labels[label]["times"].append(elapsed)
        if not success:
            labels[label]["errors"] += 1

    recap = []

    total_samples_all = 0
    all_times = []
    total_errors_all = 0

    for label, data in sorted(labels.items(), key=lambda x: x[0]):
        times = data["times"]
        errors = data["errors"]
        samples = len(times)
        if samples == 0:
            continue

        avg = statistics.mean(times)
        mn = min(times)
        mx = max(times)
        std_dev = statistics.pstdev(times) if samples > 1 else 0.0
        p90 = percentile(times, 90)
        p95 = percentile(times, 95)
        p99 = percentile(times, 99)
        err_pct = (errors / samples * 100.0) if samples else 0.0

        recap.append({
            "Label": label,
            "Samples": samples,
            "Average (ms)": round(avg, 2),
            "Min (ms)": mn,
            "Max (ms)": mx,
            "Std Dev (ms)": round(std_dev, 2),
            "90% Line (ms)": round(p90, 2) if p90 is not None else "",
            "95% Line (ms)": round(p95, 2) if p95 is not None else "",
            "99% Line (ms)": round(p99, 2) if p99 is not None else "",
            "Error %": round(err_pct, 2),
        })

        total_samples_all += samples
        total_errors_all += errors
        all_times.extend(times)

    # Ligne TOTAL
    if all_times:
        total_avg = statistics.mean(all_times)
        total_min = min(all_times)
        total_max = max(all_times)
        total_std = statistics.pstdev(all_times) if len(all_times) > 1 else 0.0
        total_err_pct = (total_errors_all / total_samples_all * 100.0) if total_samples_all else 0.0
        p90 = percentile(all_times, 90)
        p95 = percentile(all_times, 95)
        p99 = percentile(all_times, 99)

        recap.append({
            "Label": "TOTAL",
            "Samples": total_samples_all,
            "Average (ms)": round(total_avg, 2),
            "Min (ms)": total_min,
            "Max (ms)": total_max,
            "Std Dev (ms)": round(total_std, 2),
            "90% Line (ms)": round(p90, 2) if p90 is not None else "",
            "95% Line (ms)": round(p95, 2) if p95 is not None else "",
            "99% Line (ms)": round(p99, 2) if p99 is not None else "",
            "Error %": round(total_err_pct, 2),
        })

    logging.info("  -> %d lignes dans le tableau récap (y compris TOTAL)", len(recap))
    return recap


# --------------------------------------------------------
# Excel
# --------------------------------------------------------
def sanitize_sheet_name(name: str) -> str:
    name = re.sub(r'[:\\/?*\[\]]', "_", name)
    if len(name) > 31:
        name = name[:31]
    if not name:
        name = "Sheet"
    return name


def write_excel(output_file: str,
                scenarios_data: dict,
                scenarios_users: list,
                rt_matrix: dict,
                err_matrix: dict):
    logging.info("Création du fichier Excel : %s", output_file)
    workbook = xlsxwriter.Workbook(output_file)

    header_fmt = workbook.add_format({
        "bold": True,
        "bg_color": "#D9D9D9",
        "border": 1
    })
    cell_fmt = workbook.add_format({"border": 1})
    num_fmt = workbook.add_format({"border": 1, "num_format": "0.00"})
    int_fmt = workbook.add_format({"border": 1, "num_format": "0"})  # entier

    # --- Feuilles par scénario (détail) ---
    for sheet_name_raw, rows in scenarios_data.items():
        sheet_name = sanitize_sheet_name(sheet_name_raw)
        logging.info("  -> Création de la feuille : %s", sheet_name)

        ws = workbook.add_worksheet(sheet_name)

        if not rows:
            logging.warning("    (Aucune donnée pour ce scénario)")
            continue

        headers = [
            "Label",
            "Samples",
            "Average (ms)",
            "Min (ms)",
            "Max (ms)",
            "Std Dev (ms)",
            "90% Line (ms)",
            "95% Line (ms)",
            "99% Line (ms)",
            "Error %",
        ]

        for col, h in enumerate(headers):
            ws.write(0, col, h, header_fmt)

        for row_idx, row in enumerate(rows, start=1):
            for col_idx, h in enumerate(headers):
                val = row.get(h, "")
                if isinstance(val, (int, float)):
                    ws.write(row_idx, col_idx, val, num_fmt)
                else:
                    ws.write(row_idx, col_idx, str(val), cell_fmt)

        ws.set_column(0, 0, 40)
        ws.set_column(1, len(headers) - 1, 16)

    # --- Onglet Data Time Response Time (Scenario / API / Response Time) ---
    ws_rt = workbook.add_worksheet("Data Time Response Time")
    ws_rt.write(0, 0, "Scenario", header_fmt)
    ws_rt.write(0, 1, "API", header_fmt)
    ws_rt.write(0, 2, "Response Time (ms)", header_fmt)

    row_idx = 1
    for users in sorted(scenarios_users):
        start_row = row_idx
        for label in LABEL_ORDER:
            val = rt_matrix.get(label, {}).get(users, None)
            if val is None:
                continue
            # API
            ws_rt.write(row_idx, 1, label, cell_fmt)
            # Response time entier
            ws_rt.write(row_idx, 2, int(round(val)), int_fmt)
            row_idx += 1
        end_row = row_idx - 1
        if end_row >= start_row:
            # fusion de la colonne Scenario pour ce bloc
            ws_rt.merge_range(start_row, 0, end_row, 0, users, int_fmt)

    ws_rt.set_column(0, 0, 12)   # Scenario
    ws_rt.set_column(1, 1, 20)   # API
    ws_rt.set_column(2, 2, 20)   # Response Time

    # --- Onglet Data Error Rate (Scenario / API / Error Rate) ---
    ws_err = workbook.add_worksheet("Data Error Rate")
    ws_err.write(0, 0, "Scenario", header_fmt)
    ws_err.write(0, 1, "API", header_fmt)
    ws_err.write(0, 2, "Error Rate (%)", header_fmt)

    row_idx = 1
    for users in sorted(scenarios_users):
        start_row = row_idx
        for label in LABEL_ORDER:
            val = err_matrix.get(label, {}).get(users, None)
            if val is None:
                continue
            ws_err.write(row_idx, 1, label, cell_fmt)

            # formattage ErrorRate :
            # - entier si possible
            # - sinon 2 décimales avec un point (écrit comme texte)
            if abs(val - round(val)) < 1e-9:
                s = str(int(round(val)))
            else:
                s = f"{val:.2f}"  # Python utilise déjà le point comme séparateur
            ws_err.write(row_idx, 2, s, cell_fmt)

            row_idx += 1
        end_row = row_idx - 1
        if end_row >= start_row:
            ws_err.merge_range(start_row, 0, end_row, 0, users, int_fmt)

    ws_err.set_column(0, 0, 12)
    ws_err.set_column(1, 1, 20)
    ws_err.set_column(2, 2, 20)

    workbook.close()
    logging.info("Fichier Excel finalisé.")


# --------------------------------------------------------
# Word helpers
# --------------------------------------------------------
def insert_table_after_paragraph(paragraph, rows, cols):
    """
    Insère un tableau juste après un paragraphe.
    Retourne l'objet Table python-docx.
    """
    tbl = OxmlElement('w:tbl')
    paragraph._p.addnext(tbl)
    table = Table(tbl, paragraph._parent)

    # créer les lignes/colonnes
    for _ in range(rows):
        tr = OxmlElement('w:tr')
        tbl.append(tr)
        for _ in range(cols):
            tc = OxmlElement('w:tc')
            tr.append(tc)
            p = OxmlElement('w:p')
            tc.append(p)

    return table


def generate_word_report(template_path, output_path,
                         scenarios_users, scenario_recaps, scenario_rows):
    """
    Modifie le template Word :
      - remplit les dates d'exécution dans les tableaux "Execution date"
      - insère un tableau Response time sous chaque titre "Response time"
    """
    if not template_path:
        logging.warning("DOC_TEMPLATE non défini, génération Word ignorée.")
        return

    if not os.path.isfile(template_path):
        logging.error("DOC_TEMPLATE n'existe pas : %s", template_path)
        return

    logging.info("Ouverture du template Word : %s", template_path)
    doc = Document(template_path)

    # 1) Remplir les dates d'exécution dans les lignes "Execution date"
    exec_strings = []
    for users in sorted(scenarios_users):
        rows = scenario_rows.get(users, [])
        exec_strings.append(compute_execution_range_string(rows))

    exec_cells = []
    for table in doc.tables:
        for row in table.rows:
            if row.cells and row.cells[0].text.strip().lower().startswith("execution date"):
                exec_cells.append(row.cells[1])

    for i, cell in enumerate(exec_cells):
        if i < len(exec_strings):
            cell.text = exec_strings[i]

    logging.info("Dates d'exécution mises à jour pour %d scénarios.", len(exec_strings))

    # 2) Insérer les tableaux Response time sous chaque titre "Response time"
    rt_paragraphs = [p for p in doc.paragraphs
                     if p.text.strip().lower().startswith("response time")]

    headers = [
        "Label",
        "Samples",
        "Average (ms)",
        "Min (ms)",
        "Max (ms)",
        "Std Dev (ms)",
        "90% Line (ms)",
        "95% Line (ms)",
        "99% Line (ms)",
        "Error %",
    ]

    for idx, users in enumerate(sorted(scenarios_users)):
        if idx >= len(rt_paragraphs):
            break
        para = rt_paragraphs[idx]
        recap = scenario_recaps.get(users)
        if not recap:
            continue

        logging.info("Insertion du tableau Response time pour le scénario %d users", users)
        table = insert_table_after_paragraph(para, rows=len(recap) + 1, cols=len(headers))

        # remplir en-têtes
        for col, h in enumerate(headers):
            table.rows[0].cells[col].text = h

        # remplir données
        for r_idx, r in enumerate(recap, start=1):
            cells = table.rows[r_idx].cells
            cells[0].text = r["Label"]
            cells[1].text = str(r["Samples"])
            cells[2].text = str(r["Average (ms)"])
            cells[3].text = str(r["Min (ms)"])
            cells[4].text = str(r["Max (ms)"])
            cells[5].text = str(r["Std Dev (ms)"])
            cells[6].text = str(r["90% Line (ms)"])
            cells[7].text = str(r["95% Line (ms)"])
            cells[8].text = str(r["99% Line (ms)"])
            cells[9].text = str(r["Error %"])

    doc.save(output_path)
    logging.info("Document Word généré : %s", output_path)


# --------------------------------------------------------
# MAIN
# --------------------------------------------------------
def main():
    try:
        results_folder, output_file, doc_template, doc_output = load_env()
        files = find_scenario_files(results_folder)

        scenarios_data = {}
        scenarios_users = []
        rt_matrix = defaultdict(dict)   # label -> {users: avg}
        err_matrix = defaultdict(dict)  # label -> {users: error%}
        scenario_rows = {}              # users -> rows CSV
        scenario_recaps_by_users = {}   # users -> recap

        for f in files:
            logging.info("--------------------------------------------------")
            logging.info("Traitement du fichier scénario : %s", f)

            users = extract_users_from_filename(f)
            if users not in scenarios_users:
                scenarios_users.append(users)

            rows = read_jmeter_csv(f)
            recap = compute_recap(rows)

            base_name = os.path.splitext(os.path.basename(f))[0]
            scenarios_data[base_name] = recap
            scenario_rows[users] = rows
            scenario_recaps_by_users[users] = recap

            # alimenter les matrices pour les onglets Data*
            for r in recap:
                if r["Label"] == "TOTAL":
                    continue
                label = r["Label"]
                rt_matrix[label][users] = r["Average (ms)"]
                err_matrix[label][users] = r["Error %"]

        write_excel(output_file, scenarios_data, scenarios_users, rt_matrix, err_matrix)

        # Génération du Word si les chemins sont définis
        if doc_template and doc_output:
            generate_word_report(doc_template, doc_output,
                                 scenarios_users, scenario_recaps_by_users, scenario_rows)
        else:
            logging.info("DOC_TEMPLATE ou DOC_OUTPUT non défini, Word ignoré.")

        logging.info("Terminé ✅")

    except Exception as e:
        logging.exception("❌ Erreur lors de l'exécution du script : %s", e)


if __name__ == "__main__":
    main()
